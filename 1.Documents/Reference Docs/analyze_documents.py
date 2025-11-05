import openpyxl
from docx import Document
import json

def analyze_excel(filename):
    """Analyze Excel file structure and content"""
    wb = openpyxl.load_workbook(filename, data_only=True)

    result = {
        'filename': filename,
        'sheets': []
    }

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_data = {
            'name': sheet_name,
            'rows': sheet.max_row,
            'columns': sheet.max_column,
            'data_sample': []
        }

        # Get first 20 rows to understand structure
        for i, row in enumerate(sheet.iter_rows(values_only=True), 1):
            if i <= 20:
                sheet_data['data_sample'].append({
                    'row': i,
                    'values': row
                })

        result['sheets'].append(sheet_data)

    return result

def analyze_word(filename):
    """Analyze Word document structure and content"""
    doc = Document(filename)

    result = {
        'filename': filename,
        'paragraphs_count': len(doc.paragraphs),
        'tables_count': len(doc.tables),
        'paragraphs': [],
        'tables': []
    }

    # Get first 30 paragraphs
    for i, para in enumerate(doc.paragraphs[:30], 1):
        if para.text.strip():
            result['paragraphs'].append({
                'index': i,
                'text': para.text,
                'style': para.style.name
            })

    # Get table structures
    for i, table in enumerate(doc.tables, 1):
        table_data = {
            'table_number': i,
            'rows': len(table.rows),
            'columns': len(table.columns),
            'sample_data': []
        }

        # Get first 10 rows of each table
        for row_idx, row in enumerate(table.rows[:10], 1):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data['sample_data'].append({
                'row': row_idx,
                'cells': row_data
            })

        result['tables'].append(table_data)

    return result

# Analyze the files
excel_file = "SUMMARY AS AT END OF SEPTEMBER 2025 Final.xlsx"
word_file = "KTDA Factory ICT Report -  TBESONIK  MONTHLY - September 2025.doc"

print("="*80)
print("ANALYZING EXCEL FILE")
print("="*80)
excel_analysis = analyze_excel(excel_file)
print(json.dumps(excel_analysis, indent=2, default=str))

print("\n" + "="*80)
print("ANALYZING WORD DOCUMENT")
print("="*80)
word_analysis = analyze_word(word_file)
print(json.dumps(word_analysis, indent=2, default=str))
